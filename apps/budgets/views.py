import threading
from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse, JsonResponse
from django.db.models import Q
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import PasswordChangeForm
from django.contrib.auth import update_session_auth_hash, login
from django.contrib import messages
from django.contrib.auth.models import User
from .models import InfosBudget, GroupeArticle, SousLigneArticle, SectionBudgetaire
from .forms import SousLigneArticleForm, InfosBudgetForm, InscriptionOperateurForm


def _envoyer_email_async(subject, message, recipient_list):
    """Envoie un email dans un thread séparé pour ne pas bloquer la requête."""
    from django.core.mail import send_mail
    from django.conf import settings

    def _send():
        try:
            send_mail(
                subject=subject,
                message=message,
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=recipient_list,
                fail_silently=True,
            )
        except Exception as e:
            print(f"Erreur envoi email: {e}")

    threading.Thread(target=_send, daemon=True).start()

def inscription(request):
    """Inscription d'un nouvel opérateur (utilisateur sans droits admin)"""
    if request.user.is_authenticated:
        return redirect('budgets:dashboard')

    if request.method == 'POST':
        form = InscriptionOperateurForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.is_staff = False
            user.is_superuser = False
            user.save()
            login(request, user)
            messages.success(request, f"Bienvenue {user.get_full_name()} ! Votre compte opérateur a été créé avec succès.")
            return redirect('budgets:dashboard')
    else:
        form = InscriptionOperateurForm()

    return render(request, 'registration/inscription.html', {'form': form})


@login_required
def budget_detail(request, uuid):
    budget = get_object_or_404(InfosBudget, uuid=uuid)

    # Un opérateur ne peut voir que ses propres budgets
    if not (request.user.is_staff or request.user.is_superuser):
        if budget.created_by != request.user:
            messages.error(request, "Vous n'avez pas accès à ce budget.")
            return redirect('budgets:dashboard')

    # S'assurer que la structure budgétaire est initialisée
    budget.initialiser_structure()
    # Forcer le rafraîchissement des données depuis la base
    budget.refresh_from_db()
    return render(request, 'budgets/budget.html', {'budget': budget})

@login_required
def afficher_formulaire_ligne(request, groupe_id):
    """ Renvoie la ligne de formulaire <tr> """
    groupe = get_object_or_404(GroupeArticle, pk=groupe_id)
    return render(request, 'budgets/subline_form.html', {
        'groupe': groupe
    })

@login_required
def sauvegarder_ligne(request, groupe_id):
    """ Sauvegarde et demande le rafraîchissement des totaux """
    groupe = get_object_or_404(GroupeArticle, pk=groupe_id)

    # Vérifier que le budget est modifiable (statut + appel actif)
    budget = groupe.ligne.section.budget_parent
    if not budget.peut_etre_modifie():
        return HttpResponse("Ce budget ne peut pas être modifié.", status=403)

    # L'admin ne peut pas ajouter de lignes
    if request.user.is_staff or request.user.is_superuser:
        return HttpResponse("Les administrateurs ne peuvent pas modifier les budgets.", status=403)

    if request.method == "POST":
        try:
            article = SousLigneArticle.objects.create(
                groupe=groupe,
                designation=request.POST.get('designation', ''),
                unite=request.POST.get('unite', ''),
                quantite=float(request.POST.get('quantite', 1)),
                prix_unitaire=float(request.POST.get('prix_unitaire', 0)),
                co_financement=float(request.POST.get('co_financement', 0))
            )

            response = render(request, 'partials/subline_row.html', {'article': article, 'budget': budget})
            response['HX-Trigger'] = 'refreshTotals'
            return response
        except (ValueError, TypeError) as e:
            return HttpResponse(f"Erreur de données: {str(e)}", status=400)

    return HttpResponse("Méthode non autorisée", status=405)

@login_required
def supprimer_article(request, article_id):
    """Supprime un article via HTMX"""
    article = get_object_or_404(SousLigneArticle, pk=article_id)

    # Vérifier que le budget est modifiable
    budget = article.groupe.ligne.section.budget_parent
    if not budget.peut_etre_modifie():
        return HttpResponse("Ce budget ne peut pas être modifié.", status=403)

    if request.user.is_staff or request.user.is_superuser:
        return HttpResponse("Les administrateurs ne peuvent pas modifier les budgets.", status=403)

    article.delete()
    
    # Retourne une réponse vide avec le trigger pour rafraîchir les totaux
    response = HttpResponse("")
    response['HX-Trigger'] = 'refreshTotals'
    return response

@login_required
def get_synthese(request, uuid):
    budget = get_object_or_404(InfosBudget, uuid=uuid)
    budget.calculer_synthese()
    return render(request, 'budgets/synthese_header.html', {'budget': budget})

@login_required
def get_budget_statut(request, uuid):
    """Retourne le statut actuel du budget en JSON (pour le polling temps réel)"""
    budget = get_object_or_404(InfosBudget, uuid=uuid)
    return JsonResponse({
        'statut': budget.statut,
        'statut_display': budget.get_statut_display(),
        'motif': budget.motif_demande_modification or '',
    })

@login_required
def get_pourcentage_a1(request, uuid):
    """Retourne le pourcentage A1 et les infos de validation en JSON"""
    budget = get_object_or_404(InfosBudget, uuid=uuid)
    budget.calculer_synthese()
    budget.refresh_from_db()

    return JsonResponse({
        'pourcentage_a1': float(budget.pourcentage_a1),
        'cout_total_global': float(budget.cout_total_global),
        'depasse_30': float(budget.pourcentage_a1) > 30,
    })

@login_required
def get_appel_statut(request):
    """Retourne le nombre d'appels actifs (pour le polling temps réel)"""
    from apps.projects.models import AppelAProjet
    from django.utils import timezone as tz
    now = tz.now()
    count = AppelAProjet.objects.filter(
        date_debut__lte=now, date_fin__gte=now
    ).count()
    return JsonResponse({
        'count': count,
    })

@login_required
def budget_dashboard(request):
    query = request.GET.get('q')

    is_admin = request.user.is_staff or request.user.is_superuser

    if is_admin:
        # Admin : voit tous les budgets, avec ou sans recherche
        if query:
            budgets = InfosBudget.objects.filter(
                Q(titre_projet__icontains=query) |
                Q(operateur__icontains=query) |
                Q(filiere__icontains=query) |
                Q(localite__nom__icontains=query) |
                Q(metier__nom__icontains=query)
            ).order_by('-id')
        else:
            budgets = InfosBudget.objects.all().order_by('-id')
    else:
        # Opérateur : voit ses budgets UNIQUEMENT s'il fait une recherche
        if query:
            budgets = InfosBudget.objects.filter(
                Q(created_by=request.user),
                Q(titre_projet__icontains=query) |
                Q(operateur__icontains=query) |
                Q(filiere__icontains=query) |
                Q(localite__nom__icontains=query) |
                Q(metier__nom__icontains=query)
            ).order_by('-id')
        else:
            budgets = InfosBudget.objects.none()

    # Récupérer tous les appels à projet actifs
    from apps.projects.models import AppelAProjet
    from django.utils import timezone as tz
    now = tz.now()
    appels_actifs = AppelAProjet.objects.filter(
        date_debut__lte=now, date_fin__gte=now
    )

    return render(request, 'budgets/dashboard.html', {
        'budgets': budgets,
        'query': query,
        'is_admin': is_admin,
        'appels_actifs': appels_actifs,
    })

@login_required
def creer_budget(request):
    from apps.projects.models import AppelAProjet
    from django.utils import timezone

    # L'admin ne peut pas créer de budget
    if request.user.is_staff or request.user.is_superuser:
        messages.error(request, "Les administrateurs ne sont pas autorisés à créer des budgets.")
        return redirect('budgets:dashboard')

    # Récupérer tous les appels à projet actifs
    now = timezone.now()
    appels_actifs = AppelAProjet.objects.filter(
        date_debut__lte=now, date_fin__gte=now
    )
    if not appels_actifs.exists():
        messages.error(request, "Aucun appel à projet n'est actif actuellement. Vous ne pouvez pas créer de budget en dehors de la période d'appel.")
        return redirect('budgets:dashboard')

    if request.method == "POST":
        form = InfosBudgetForm(request.POST, appels_actifs=appels_actifs)
        if form.is_valid():
            nouveau_budget = form.save(commit=False)
            nouveau_budget.created_by = request.user
            nouveau_budget.save()
            return redirect('budgets:budget_detail', uuid=nouveau_budget.uuid)
    else:
        form = InfosBudgetForm(appels_actifs=appels_actifs)

    return render(request, 'budgets/creer_budget.html', {'form': form, 'appels_actifs': appels_actifs})

@login_required
def modifier_budget(request, uuid):
    """Modifie les infos d'un budget - interdit aux admins"""
    # L'admin ne peut pas modifier les budgets
    if request.user.is_staff or request.user.is_superuser:
        messages.error(request, "Les administrateurs ne sont pas autorisés à modifier les budgets.")
        return redirect('budgets:budget_detail', uuid=uuid)

    budget = get_object_or_404(InfosBudget, uuid=uuid)

    # Vérifier que l'opérateur est le créateur
    if budget.created_by != request.user:
        messages.error(request, "Vous n'avez pas accès à ce budget.")
        return redirect('budgets:dashboard')

    # Vérifier que le budget est modifiable (statut + appel actif)
    if not budget.peut_etre_modifie():
        messages.error(request, "Ce budget ne peut pas être modifié dans son état actuel.")
        return redirect('budgets:budget_detail', uuid=uuid)

    if request.method == "POST":
        form = InfosBudgetForm(request.POST, instance=budget)
        if form.is_valid():
            form.save()
            return redirect('budgets:budget_detail', uuid=budget.uuid)
    else:
        form = InfosBudgetForm(instance=budget)

    return render(request, 'budgets/modifier_budget.html', {'form': form, 'budget': budget})

@login_required
def supprimer_budget(request, uuid):
    """Suppression désactivée sur la plateforme. Utiliser l'admin Django."""
    messages.error(request, "La suppression de budgets n'est pas autorisée depuis la plateforme. Contactez l'administrateur Django.")
    return redirect('budgets:dashboard')

@login_required
def get_total_section(request, section_id):
    section = get_object_or_404(SectionBudgetaire, id=section_id)
    return render(request, 'budgets/partials/section_total.html', {
        'section': section
    })

# ========================= EXPORTS =========================

@login_required
def export_excel(request, uuid):
    """Exporte un budget en format Excel (réservé aux admins)"""
    # Vérifier que l'utilisateur est admin
    if not (request.user.is_staff or request.user.is_superuser):
        return HttpResponse("Accès non autorisé - Seuls les administrateurs peuvent exporter", status=403)

    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    budget = get_object_or_404(InfosBudget, uuid=uuid)
    budget.calculer_synthese()

    # Créer un classeur Excel
    wb = Workbook()
    ws = wb.active
    ws.title = "Budget"

    # Styles
    header_fill = PatternFill(start_color="667eea", end_color="667eea", fill_type="solid")
    section_fill = PatternFill(start_color="343a40", end_color="343a40", fill_type="solid")
    ligne_fill = PatternFill(start_color="6c757d", end_color="6c757d", fill_type="solid")
    groupe_fill = PatternFill(start_color="0dcaf0", end_color="0dcaf0", fill_type="solid")

    header_font = Font(bold=True, color="FFFFFF", size=12)
    section_font = Font(bold=True, color="FFFFFF", size=11)
    ligne_font = Font(bold=True, color="FFFFFF", size=10)
    groupe_font = Font(bold=True, color="000000", size=10)

    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    center_alignment = Alignment(horizontal='center', vertical='center')
    right_alignment = Alignment(horizontal='right', vertical='center')

    # En-tête du document
    ws.merge_cells('A1:I1')
    cell = ws['A1']
    cell.value = f"BUDGET - {budget.titre_projet}"
    cell.font = Font(bold=True, size=14, color="667eea")
    cell.alignment = center_alignment

    # Informations générales
    row = 3
    ws[f'A{row}'] = "Opérateur:"
    ws[f'B{row}'] = str(budget.operateur)
    ws[f'E{row}'] = "Filière:"
    ws[f'F{row}'] = budget.filiere
    row += 1
    ws[f'A{row}'] = "Apprenants:"
    ws[f'B{row}'] = budget.total_apprenants
    ws[f'E{row}'] = "Sessions:"
    ws[f'F{row}'] = budget.nombre_sessions

    # Synthèse
    row += 2
    ws.merge_cells(f'A{row}:I{row}')
    cell = ws[f'A{row}']
    cell.value = "SYNTHÈSE DU BUDGET"
    cell.font = header_font
    cell.fill = header_fill
    cell.alignment = center_alignment

    row += 1
    synthese_data = [
        ["COÛT TOTAL GLOBAL", f"{budget.cout_total_global:,.0f} FCFA"],
        ["BUDGET DEMANDÉ", f"{budget.budget_demande_global:,.0f} FCFA"],
        ["CO-FINANCEMENT", f"{budget.co_financement_global:,.0f} FCFA"],
        ["COÛT / APPRENANT", f"{budget.cout_par_apprenant:,.0f} FCFA"],
    ]

    for label, value in synthese_data:
        ws[f'A{row}'] = label
        ws[f'B{row}'] = value
        ws[f'A{row}'].font = Font(bold=True)
        row += 1

    # En-têtes des colonnes
    row += 1
    headers = ["Code", "Désignation", "Unité", "Quantité", "Prix Unit.", "Coût Total", "Co-financement", "Budget demandé", ""]
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=row, column=col_num)
        cell.value = header
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_alignment
        cell.border = thin_border

    row += 1

    # Parcourir les sections
    for section in budget.sections.all():
        # Ligne de section
        ws.merge_cells(f'A{row}:E{row}')
        cell = ws[f'A{row}']
        cell.value = f"{section.code} - {section.libelle}"
        cell.font = section_font
        cell.fill = section_fill

        ws[f'F{row}'] = section.cout_total_section
        ws[f'G{row}'] = section.co_financement_section
        ws[f'H{row}'] = section.budget_demande_section

        for col in ['F', 'G', 'H']:
            ws[f'{col}{row}'].font = section_font
            ws[f'{col}{row}'].fill = section_fill
            ws[f'{col}{row}'].alignment = right_alignment
            ws[f'{col}{row}'].number_format = '#,##0'

        row += 1

        # Parcourir les lignes
        for ligne in section.lignes.all():
            ws.merge_cells(f'A{row}:E{row}')
            cell = ws[f'A{row}']
            cell.value = f"  {ligne.code} - {ligne.libelle}"
            cell.font = ligne_font
            cell.fill = ligne_fill

            ws[f'F{row}'] = ligne.cout_total_ligne
            ws[f'G{row}'] = ligne.co_financement_ligne
            ws[f'H{row}'] = ligne.budget_demande_ligne

            for col in ['F', 'G', 'H']:
                ws[f'{col}{row}'].font = ligne_font
                ws[f'{col}{row}'].fill = ligne_fill
                ws[f'{col}{row}'].alignment = right_alignment
                ws[f'{col}{row}'].number_format = '#,##0'

            row += 1

            # Parcourir les groupes
            for groupe in ligne.groupes.all():
                ws.merge_cells(f'A{row}:H{row}')
                cell = ws[f'A{row}']
                cell.value = f"    {groupe.libelle}"
                cell.font = groupe_font
                cell.fill = groupe_fill
                row += 1

                # Parcourir les articles
                for article in groupe.articles.all():
                    ws[f'A{row}'] = ""
                    ws[f'B{row}'] = article.designation
                    ws[f'C{row}'] = article.unite
                    ws[f'D{row}'] = article.quantite
                    ws[f'E{row}'] = article.prix_unitaire
                    ws[f'F{row}'] = article.cout_total
                    ws[f'G{row}'] = article.co_financement
                    ws[f'H{row}'] = article.budget_demande_article

                    for col in ['D', 'E', 'F', 'G', 'H']:
                        ws[f'{col}{row}'].alignment = right_alignment
                        ws[f'{col}{row}'].number_format = '#,##0.00'

                    row += 1

    # Ajuster la largeur des colonnes
    column_widths = [10, 40, 12, 12, 15, 15, 18, 15, 5]
    for i, width in enumerate(column_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = width

    # Créer la réponse HTTP
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    filename = f"Budget_{budget.titre_projet.replace(' ', '_')}_{budget.uuid}.xlsx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    wb.save(response)
    return response


@login_required
def export_pdf(request, uuid):
    """Exporte un budget en format PDF (réservé aux admins)"""
    # Vérifier que l'utilisateur est admin
    if not (request.user.is_staff or request.user.is_superuser):
        return HttpResponse("Accès non autorisé - Seuls les administrateurs peuvent exporter", status=403)

    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from io import BytesIO

    budget = get_object_or_404(InfosBudget, uuid=uuid)
    budget.calculer_synthese()

    # Créer le buffer
    buffer = BytesIO()

    # Créer le document en format paysage pour avoir plus d'espace
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        rightMargin=1*cm,
        leftMargin=1*cm,
        topMargin=1*cm,
        bottomMargin=1*cm
    )

    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#667eea'),
        alignment=TA_CENTER,
        spaceAfter=12
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=12,
        textColor=colors.white,
        alignment=TA_CENTER,
        spaceAfter=6
    )

    # Contenu
    elements = []

    # Titre
    elements.append(Paragraph(f"BUDGET - {budget.titre_projet}", title_style))
    elements.append(Spacer(1, 0.3*cm))

    # Informations générales
    info_data = [
        ["Opérateur:", str(budget.operateur), "Filière:", budget.filiere],
        ["Apprenants:", str(budget.total_apprenants), "Sessions:", str(budget.nombre_sessions)],
    ]

    info_table = Table(info_data, colWidths=[4*cm, 8*cm, 3*cm, 8*cm])
    info_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 0.5*cm))

    # Synthèse
    elements.append(Paragraph("SYNTHÈSE DU BUDGET", heading_style))
    synthese_data = [
        ["COÛT TOTAL GLOBAL", f"{budget.cout_total_global:,.0f} FCFA"],
        ["BUDGET DEMANDÉ", f"{budget.budget_demande_global:,.0f} FCFA"],
        ["CO-FINANCEMENT", f"{budget.co_financement_global:,.0f} FCFA"],
        ["COÛT / APPRENANT", f"{budget.cout_par_apprenant:,.0f} FCFA"],
    ]

    synthese_table = Table(synthese_data, colWidths=[8*cm, 8*cm])
    synthese_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#667eea')),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.white),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ROWBACKGROUNDS', (1, 0), (1, -1), [colors.white, colors.lightgrey]),
    ]))
    elements.append(synthese_table)
    elements.append(Spacer(1, 0.5*cm))

    # Données du budget
    data = [
        ["Code", "Désignation", "Unité", "Qté", "Prix Unit.", "Coût Total", "Co-fin.", "Budget demandé"]
    ]

    for section in budget.sections.all():
        data.append([
            section.code,
            section.libelle,
            "",
            "",
            "",
            f"{section.cout_total_section:,.0f}",
            f"{section.co_financement_section:,.0f}",
            f"{section.budget_demande_section:,.0f}"
        ])

        for ligne in section.lignes.all():
            data.append([
                f"  {ligne.code}",
                ligne.libelle,
                "",
                "",
                "",
                f"{ligne.cout_total_ligne:,.0f}",
                f"{ligne.co_financement_ligne:,.0f}",
                f"{ligne.budget_demande_ligne:,.0f}"
            ])

            for groupe in ligne.groupes.all():
                data.append([
                    "",
                    f"    {groupe.libelle}",
                    "",
                    "",
                    "",
                    "",
                    "",
                    ""
                ])

                for article in groupe.articles.all():
                    data.append([
                        "",
                        f"      {article.designation}",
                        article.unite or "",
                        f"{article.quantite:,.2f}",
                        f"{article.prix_unitaire:,.0f}",
                        f"{article.cout_total:,.0f}",
                        f"{article.co_financement:,.0f}",
                        f"{article.budget_demande_article:,.0f}"
                    ])

    # Créer le tableau
    col_widths = [2*cm, 6*cm, 2*cm, 1.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm]
    table = Table(data, colWidths=col_widths, repeatRows=1)

    # Style du tableau
    table_style = [
        # En-tête
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#667eea')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),

        # Corps
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('ALIGN', (3, 1), (-1, -1), 'RIGHT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
    ]

    table.setStyle(TableStyle(table_style))
    elements.append(table)

    # Construire le PDF
    doc.build(elements)

    # Obtenir la valeur du buffer
    pdf = buffer.getvalue()
    buffer.close()

    # Créer la réponse
    response = HttpResponse(content_type='application/pdf')
    filename = f"Budget_{budget.titre_projet.replace(' ', '_')}_{budget.uuid}.pdf"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    response.write(pdf)

    return response


@login_required
def export_word(request, uuid):
    """Exporte un budget en format Word (réservé aux admins)"""
    # Vérifier que l'utilisateur est admin
    if not (request.user.is_staff or request.user.is_superuser):
        return HttpResponse("Accès non autorisé - Seuls les administrateurs peuvent exporter", status=403)

    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    budget = get_object_or_404(InfosBudget, uuid=uuid)
    budget.calculer_synthese()

    # Créer le document
    doc = Document()

    # Titre
    title = doc.add_heading(f'BUDGET - {budget.titre_projet}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(102, 126, 234)

    # Informations générales
    doc.add_heading('Informations générales', level=1)
    info_table = doc.add_table(rows=2, cols=4)
    info_table.style = 'Light Grid Accent 1'

    cells = info_table.rows[0].cells
    cells[0].text = 'Opérateur:'
    cells[1].text = str(budget.operateur)
    cells[2].text = 'Filière:'
    cells[3].text = budget.filiere

    cells = info_table.rows[1].cells
    cells[0].text = 'Apprenants:'
    cells[1].text = str(budget.total_apprenants)
    cells[2].text = 'Sessions:'
    cells[3].text = str(budget.nombre_sessions)

    doc.add_paragraph()

    # Synthèse
    doc.add_heading('Synthèse du Budget', level=1)
    synthese_table = doc.add_table(rows=5, cols=2)
    synthese_table.style = 'Medium Shading 1 Accent 1'

    synthese_data = [
        ["Indicateur", "Montant"],
        ["COÛT TOTAL GLOBAL", f"{budget.cout_total_global:,.0f} FCFA"],
        ["BUDGET DEMANDÉ", f"{budget.budget_demande_global:,.0f} FCFA"],
        ["CO-FINANCEMENT", f"{budget.co_financement_global:,.0f} FCFA"],
        ["COÛT / APPRENANT", f"{budget.cout_par_apprenant:,.0f} FCFA"],
    ]

    for i, (label, value) in enumerate(synthese_data):
        cells = synthese_table.rows[i].cells
        cells[0].text = label
        cells[1].text = value

    doc.add_paragraph()

    # Détails du budget
    doc.add_heading('Détails du Budget', level=1)

    # En-tête du tableau
    budget_table = doc.add_table(rows=1, cols=8)
    budget_table.style = 'Light Grid Accent 1'

    headers = ["Code", "Désignation", "Unité", "Qté", "Prix Unit.", "Coût Total", "Co-fin.", "Budget demandé"]
    header_cells = budget_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Mettre en gras
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Parcourir les sections
    for section in budget.sections.all():
        row_cells = budget_table.add_row().cells
        row_cells[0].text = section.code
        row_cells[1].text = section.libelle
        row_cells[5].text = f"{section.cout_total_section:,.0f}"
        row_cells[6].text = f"{section.co_financement_section:,.0f}"
        row_cells[7].text = f"{section.budget_demande_section:,.0f}"

        # Style pour les sections
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Parcourir les lignes
        for ligne in section.lignes.all():
            row_cells = budget_table.add_row().cells
            row_cells[0].text = f"  {ligne.code}"
            row_cells[1].text = ligne.libelle
            row_cells[5].text = f"{ligne.cout_total_ligne:,.0f}"
            row_cells[6].text = f"{ligne.co_financement_ligne:,.0f}"
            row_cells[7].text = f"{ligne.budget_demande_ligne:,.0f}"

            # Parcourir les groupes
            for groupe in ligne.groupes.all():
                row_cells = budget_table.add_row().cells
                row_cells[1].text = f"    {groupe.libelle}"

                # Parcourir les articles
                for article in groupe.articles.all():
                    row_cells = budget_table.add_row().cells
                    row_cells[1].text = f"      {article.designation}"
                    row_cells[2].text = article.unite or ""
                    row_cells[3].text = f"{article.quantite:,.2f}"
                    row_cells[4].text = f"{article.prix_unitaire:,.0f}"
                    row_cells[5].text = f"{article.cout_total:,.0f}"
                    row_cells[6].text = f"{article.co_financement:,.0f}"
                    row_cells[7].text = f"{article.budget_demande_article:,.0f}"

    # Sauvegarder dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Créer la réponse
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    filename = f"Budget_{budget.titre_projet.replace(' ', '_')}_{budget.uuid}.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    return response


# ========================= PROFIL & MOT DE PASSE =========================

@login_required
def profil(request):
    """Page de profil de l'utilisateur"""
    return render(request, 'budgets/profil.html', {
        'user': request.user
    })

@login_required
def changer_mot_de_passe(request):
    """Permet à l'utilisateur de changer son mot de passe"""
    if request.method == 'POST':
        form = PasswordChangeForm(request.user, request.POST)
        if form.is_valid():
            user = form.save()
            # Mettre à jour la session pour éviter la déconnexion
            update_session_auth_hash(request, user)
            messages.success(request, 'Votre mot de passe a été changé avec succès !')
            return redirect('budgets:profil')
        else:
            messages.error(request, 'Veuillez corriger les erreurs ci-dessous.')
    else:
        form = PasswordChangeForm(request.user)

    return render(request, 'budgets/changer_mot_de_passe.html', {
        'form': form
    })

# ========================= WORKFLOW DE VALIDATION =========================

@login_required
def soumettre_budget(request, uuid):
    """Soumet le budget pour validation par l'admin"""
    from django.utils import timezone
    budget = get_object_or_404(InfosBudget, uuid=uuid)

    # Vérifier que l'utilisateur est le créateur ou admin
    if budget.created_by != request.user and not (request.user.is_staff or request.user.is_superuser):
        messages.error(request, "Vous n'avez pas la permission de soumettre ce budget.")
        return redirect('budgets:budget_detail', uuid=uuid)

    # Vérifier que le budget est modifiable (brouillon, soumis avec appel actif, ou modification autorisée)
    if not budget.peut_etre_modifie():
        messages.error(request, "Ce budget ne peut pas être soumis dans son état actuel.")
        return redirect('budgets:budget_detail', uuid=uuid)

    # Recalculer la synthèse et vérifier la règle des 30% pour A.1
    budget.calculer_synthese()
    budget.refresh_from_db()
    valide, erreur = budget.verifier_validation_a1()
    if not valide:
        messages.error(request, erreur)
        return redirect('budgets:budget_detail', uuid=uuid)

    # Mettre à jour le statut
    budget.statut = InfosBudget.STATUT_SOUMIS
    budget.date_soumission = timezone.now()
    budget.save()

    # Envoyer un email à l'admin (en arrière-plan)
    admin_emails = [u.email for u in User.objects.filter(is_staff=True, is_active=True) if u.email]
    if admin_emails:
        _envoyer_email_async(
            subject=f'Nouveau budget soumis: {budget.titre_projet}',
            message=f"""Un nouveau budget a été soumis pour validation.

Titre du projet: {budget.titre_projet}
Opérateur: {budget.operateur}
Filière: {budget.filiere}
Soumis par: {budget.created_by.username if budget.created_by else 'Inconnu'}
Date de soumission: {timezone.now().strftime('%d/%m/%Y %H:%M')}

Coût total global: {budget.cout_total_global:,.0f} FCFA
Budget demandé: {budget.budget_demande_global:,.0f} FCFA

Veuillez vous connecter à l'application pour examiner ce budget.""",
            recipient_list=admin_emails,
        )

    messages.success(request, "Votre budget a été soumis avec succès. Vous recevrez une notification une fois qu'il aura été examiné.")
    return redirect('budgets:budget_detail', uuid=uuid)

@login_required
def demander_modification(request, uuid):
    """Admin demande à l'opérateur de modifier son budget (avec motif + email)"""
    from django.utils import timezone
    # Réservé aux admins
    if not (request.user.is_staff or request.user.is_superuser):
        messages.error(request, "Accès non autorisé.")
        return redirect('budgets:dashboard')

    budget = get_object_or_404(InfosBudget, uuid=uuid)

    # Vérifier que le budget est soumis
    if budget.statut != InfosBudget.STATUT_SOUMIS:
        messages.error(request, "Seuls les budgets soumis peuvent faire l'objet d'une demande de modification.")
        return redirect('budgets:budget_detail', uuid=uuid)

    if request.method == 'POST':
        motif = request.POST.get('motif', '').strip()
        if not motif:
            messages.error(request, "Veuillez indiquer le motif de la demande de modification.")
            return render(request, 'budgets/demander_modification.html', {'budget': budget})

        # Passer directement en modification autorisée
        budget.statut = InfosBudget.STATUT_MODIFICATION_AUTORISEE
        budget.motif_demande_modification = motif
        budget.date_demande_modification = timezone.now()
        budget.date_autorisation_modification = timezone.now()
        budget.save()

        # Envoyer un email à l'opérateur (en arrière-plan)
        budget_url = request.build_absolute_uri(f'/budgets/{budget.uuid}/')
        if budget.created_by and budget.created_by.email:
            _envoyer_email_async(
                subject=f'[MODIFICATION DEMANDÉE] {budget.titre_projet}',
                message=f"""Bonjour {budget.created_by.get_full_name() or budget.created_by.username},

L'administrateur vous demande de modifier votre budget "{budget.titre_projet}".

Motif de la demande :
{motif}

Vous pouvez maintenant modifier votre budget et le soumettre à nouveau.

Cliquez ici pour accéder à votre budget :
{budget_url}

Cordialement,
Système de Gestion de Budget""",
                recipient_list=[budget.created_by.email],
            )

        messages.success(request, "La demande de modification a été envoyée à l'opérateur par email. Le budget est maintenant modifiable.")
        return redirect('budgets:budget_detail', uuid=uuid)

    return render(request, 'budgets/demander_modification.html', {'budget': budget})

@login_required
def approuver_budget(request, uuid):
    """Admin approuve définitivement un budget (vue admin uniquement)"""
    from django.utils import timezone
    # Vérifier que l'utilisateur est admin
    if not (request.user.is_staff or request.user.is_superuser):
        messages.error(request, "Accès non autorisé - Seuls les administrateurs peuvent approuver des budgets.")
        return redirect('budgets:dashboard')

    budget = get_object_or_404(InfosBudget, uuid=uuid)

    # Vérifier que le budget est soumis
    if budget.statut != InfosBudget.STATUT_SOUMIS:
        messages.error(request, "Seuls les budgets soumis peuvent être approuvés.")
        return redirect('budgets:budget_detail', uuid=uuid)

    # Approuver le budget
    budget.statut = InfosBudget.STATUT_APPROUVE
    budget.date_approbation = timezone.now()
    budget.save()

    # Envoyer un email au créateur (en arrière-plan)
    if budget.created_by and budget.created_by.email:
        _envoyer_email_async(
            subject=f'Budget approuvé: {budget.titre_projet}',
            message=f"""Félicitations! Votre budget "{budget.titre_projet}" a été approuvé.

Coût total global: {budget.cout_total_global:,.0f} FCFA
Budget demandé: {budget.budget_demande_global:,.0f} FCFA
Date d'approbation: {timezone.now().strftime('%d/%m/%Y %H:%M')}

Merci pour votre soumission.""",
            recipient_list=[budget.created_by.email],
        )

    messages.success(request, "Le budget a été approuvé. L'opérateur a été notifié.")
    return redirect('budgets:budget_detail', uuid=uuid)

@login_required
def rejeter_budget(request, uuid):
    """Admin rejette définitivement un budget"""
    from django.utils import timezone
    if not (request.user.is_staff or request.user.is_superuser):
        messages.error(request, "Accès non autorisé.")
        return redirect('budgets:dashboard')

    budget = get_object_or_404(InfosBudget, uuid=uuid)

    if budget.statut != InfosBudget.STATUT_SOUMIS:
        messages.error(request, "Seuls les budgets soumis peuvent être rejetés.")
        return redirect('budgets:budget_detail', uuid=uuid)

    budget.statut = InfosBudget.STATUT_REJETE
    budget.save()

    # Envoyer un email à l'opérateur (en arrière-plan)
    if budget.created_by and budget.created_by.email:
        _envoyer_email_async(
            subject=f'Budget rejeté: {budget.titre_projet}',
            message=f"""Bonjour {budget.created_by.get_full_name() or budget.created_by.username},

Votre budget "{budget.titre_projet}" a été rejeté par l'administrateur.

Coût total global: {budget.cout_total_global:,.0f} FCFA
Budget demandé: {budget.budget_demande_global:,.0f} FCFA

Pour toute question, veuillez contacter l'administrateur.

Cordialement,
Système de Gestion de Budget""",
            recipient_list=[budget.created_by.email],
        )

    messages.success(request, "Le budget a été rejeté. L'opérateur a été notifié.")
    return redirect('budgets:budget_detail', uuid=uuid)

